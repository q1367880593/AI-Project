from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT = Path("汇报用/爱聊留资窗口项目_本周进度简报_20260720.docx")


# 配色方案（与立项文档保持一致）
BLUE = RGBColor(31, 78, 121)
DARK = RGBColor(23, 32, 42)
MUTED = RGBColor(90, 99, 113)
GREEN = RGBColor(34, 139, 87)
ORANGE = RGBColor(210, 120, 40)
LIGHT_BLUE = "E8EEF7"
LIGHT_GRAY = "F4F6F8"
BORDER = "C9D2DF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = BORDER) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def style_table(table, header=True) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(3)
                paragraph.paragraph_format.line_spacing = 1.15
                for run in paragraph.runs:
                    run.font.name = "Microsoft YaHei"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    run.font.size = Pt(9)
        if header and i == 0:
            for cell in row.cells:
                set_cell_shading(cell, LIGHT_BLUE)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = BLUE


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.color.rgb = BLUE if level <= 2 else DARK
        run.font.bold = True
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)


def add_para(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.color.rgb = DARK
        rest = text[len(bold_prefix):]
        p.add_run(rest)
    else:
        p.add_run(text)
    for run in p.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10.5)
        if run.font.color.rgb is None:
            run.font.color.rgb = DARK


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(item)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10)
        run.font.color.rgb = DARK


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(item)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10)
        run.font.color.rgb = DARK


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_border(cell, "D6DEE8")
    set_cell_margins(cell, 140, 180, 140, 180)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(10.5)
    r.font.color.rgb = BLUE
    p.add_run("\n" + body)
    for run in p.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        if run.font.size is None:
            run.font.size = Pt(10)
    doc.add_paragraph()


def set_doc_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = DARK
    normal.paragraph_format.line_spacing = 1.2
    normal.paragraph_format.space_after = Pt(6)


def add_table_from_rows(doc: Document, rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    style_table(table)
    for i, row_values in enumerate(rows):
        for j, value in enumerate(row_values):
            cell = table.cell(i, j)
            cell.text = value
            if widths:
                cell.width = Cm(widths[j])
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Microsoft YaHei"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    run.font.size = Pt(9)
    doc.add_paragraph()


def add_table_with_status(doc: Document, rows: list[list[str]], status_col: int) -> None:
    """带状态列着色的表格：已完成绿色、进行中橙色"""
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    style_table(table)
    for i, row_values in enumerate(rows):
        for j, value in enumerate(row_values):
            cell = table.cell(i, j)
            cell.text = value
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Microsoft YaHei"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    run.font.size = Pt(9)
                    if i > 0 and j == status_col:
                        if value == "已完成":
                            run.font.color.rgb = GREEN
                            run.bold = True
                        elif value == "进行中":
                            run.font.color.rgb = ORANGE
                            run.bold = True
    doc.add_paragraph()


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("爱聊留资窗口项目")
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("本周进度简报")
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(16)
    r.font.color.rgb = DARK

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("周期：2026-07-20 ~ 2026-07-26")
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(10.5)
    r.font.color.rgb = MUTED
    doc.add_paragraph()


def build() -> None:
    doc = Document()
    set_doc_styles(doc)
    add_title(doc)

    # 一、项目概述
    add_heading(doc, "一、项目概述", 1)
    add_para(
        doc,
        "爱聊留资窗口项目旨在用户与经纪人 IM 沟通过程中，基于用户浏览搜索轨迹和完整会话内容，实时判断当前是否进入适合引导留资的“留资窗口”，辅助经纪人在合适时机以自然话术引导用户留手机号，提升留资转化并降低打扰。",
    )
    add_para(
        doc,
        "一期采用 “PHP 业务服务 + Python AI 决策服务 + 模型 Gateway” 三段式架构，AI 只辅助判断，不自动向用户索要联系方式。",
    )

    # 二、整体进展概览
    add_heading(doc, "二、整体进展概览", 1)
    add_table_with_status(
        doc,
        [
            ["阶段", "状态"],
            ["数据分析", "已完成"],
            ["分数权重设计", "已完成"],
            ["用户交互设计", "已完成"],
            ["技术方案设计", "已完成"],
            ["技术方案投产", "进行中"],
            ["AI 底层能力建设", "进行中"],
        ],
        status_col=1,
    )

    # 三、本周已完成
    add_heading(doc, "三、本周已完成", 1)

    add_heading(doc, "3.1 数据分析", 2)
    add_bullets(
        doc,
        [
            "完成历史爱聊会话数据脱敏与清洗。",
            "产出 conversation_features.jsonl 及 aggregate_stats 聚合统计。",
            "识别出留资成功会话的关键行为与话术信号，作为后续规则权重校准基线。",
        ],
    )

    add_heading(doc, "3.2 分数权重设计", 2)
    add_bullets(
        doc,
        [
            "确立三项基础分模型：用户意向 35% + 需求明确 30% + 核心留资理由 35%。",
            "设计时机轮次系数：1–3 轮抑制过早索要、4–10 轮在有效沟通时提升、超期衰减。",
            "设计软风险扣分项：拒绝、反感、不匹配等。",
            "最终分公式：max(0, min(100, 基础分 × 时机系数) - 软风险扣分)，80 分阈值触发留资提示。",
        ],
    )

    add_heading(doc, "3.3 用户交互设计", 2)
    add_bullets(
        doc,
        [
            "完成 B 端提示分级展示策略：none / passive / soft / strong / warning。",
            "明确各级别展示位置与卡片样式：strong 展示提示卡，soft 输入框上方建议，passive 侧边展示，warning 不建议要电话。",
            "完成经纪人反馈采集闭环设计：采纳 / 修改 / 忽略 / 发送，回写评测体系。",
        ],
    )

    add_heading(doc, "3.4 技术方案设计", 2)
    add_bullets(
        doc,
        [
            "产出《爱聊留资智能提示技术方案 v0.1》。",
            "明确 LLM 与规则引擎职责分工：LLM 负责语义理解和多维度标签，规则引擎负责评分和最终决策。",
            "完成两个核心接口设计：行为特征抽取接口（进聊天窗口调一次）、留资窗口评估接口（每条有效消息触发）。",
            "完成可靠性设计：幂等、版本控制、主备模型、降级链、灰区复核、配置包一键回滚。",
        ],
    )

    # 四、本周进行中
    add_heading(doc, "四、本周进行中", 1)

    add_heading(doc, "4.1 技术方案投产", 2)
    add_para(doc, "基于技术方案完成模块拆解与排期，识别 10 个核心模块，纯实现工时合计 44.5 pd。")
    add_table_from_rows(
        doc,
        [
            ["端", "模块", "工时"],
            ["后端-AI", "M1 AI 服务基础架构", "6.0"],
            ["后端-AI", "M2 行为特征抽取", "4.0"],
            ["后端-AI", "M3 评估接口-输入前置", "4.0"],
            ["后端-AI", "M4 LLM 标签提取", "5.5"],
            ["后端-AI", "M5 规则引擎", "6.5"],
            ["后端-AI", "M6 灰区复核", "3.0"],
            ["后端-AI", "M7 风控与会话状态", "3.0"],
            ["后端-配置", "M8 配置包与回滚", "6.0"],
            ["后端-PHP", "M9 PHP 端改造", "4.5"],
            ["前端", "M10 B 端展示", "2.5"],
            ["合计", "", "44.5"],
        ],
        [3.5, 8.5, 3.0],
    )
    add_bullets(
        doc,
        [
            "已启动 M1 AI 服务基础架构与 M2 行为特征抽取接口开发。",
        ],
    )

    add_heading(doc, "4.2 AI 底层能力建设", 2)
    add_bullets(
        doc,
        [
            "模型 Gateway 适配：主备模型切换、超时重试、温度等配置化能力搭建中。",
            "配置快照机制：config_bundle_id 绑定模型 / Prompt / 规则 / 灰度四元组，每次请求读取快照。",
            "日志与回放基础：request_id 全链路打通，输入输出落盘支持离线回放。",
        ],
    )

    # 五、下周计划
    add_heading(doc, "五、下周计划", 1)
    add_numbered(
        doc,
        [
            "完成 M1 AI 服务基础架构，模型 Gateway 可通。",
            "完成 M2 行为特征抽取接口，支持幂等与缓存复用。",
            "启动 M3 评估接口输入前置与 M4 LLM 标签提取 Prompt 设计。",
            "启动 M8 配置包设计与 Schema 校验，为一键回滚打基础。",
            "启动 M9 PHP 端行为特征缓存改造。",
        ],
    )

    # 六、风险与待决
    add_heading(doc, "六、风险与待决", 1)
    add_table_from_rows(
        doc,
        [
            ["风险项", "说明", "应对"],
            ["Prompt 调优周期", "M4 Prompt 与 M5 规则权重校准是质量决定项，实际投入通常大于编码工时", "预留 30–50% 调参 buffer，算法 / 产品 W4 起介入"],
            ["标注集质量", "离线评测可信度依赖标注集，覆盖场景需完整", "建议下周开始攒脱敏样本，覆盖强意向 / 拒绝 / 结束 / 方言等"],
            ["搜推 MCP 接入", "聊天中推荐房源卡片依赖搜推后端 MCP 能力", "已完成数据结构定义，待搜推侧确认 tool schema 后联调"],
            ["上线节奏", "A/B 周期取决于业务流量，可能拉长", "影子模式先行，小流量灰度逐步放量"],
        ],
        [3.0, 6.0, 6.0],
    )

    # 七、下一步需要协调
    add_heading(doc, "七、下一步需要协调", 1)
    add_numbered(
        doc,
        [
            "算法 / 产品侧确认 Prompt 评测样本与标注规范。",
            "搜推后端侧确认 search_property MCP tool schema 与排期。",
            "配置中心侧确认 config_bundle_id 指针原子切换方案。",
        ],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"已生成: {OUT}")


if __name__ == "__main__":
    build()
