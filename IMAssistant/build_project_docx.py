from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT = Path("AI留资窗口识别项目立项文档.docx")


BLUE = RGBColor(31, 78, 121)
DARK = RGBColor(23, 32, 42)
MUTED = RGBColor(90, 99, 113)
LIGHT_BLUE = "E8EEF7"
LIGHT_GRAY = "F4F6F8"
BORDER = "C9D2DF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = BORDER) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def style_table(table, header=True) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(3)
                paragraph.paragraph_format.line_spacing = 1.15
                for run in paragraph.runs:
                    run.font.name = "Microsoft YaHei"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    run.font.size = Pt(9)
        if header and i == 0:
            for cell in row.cells:
                set_cell_shading(cell, LIGHT_BLUE)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = BLUE


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.color.rgb = BLUE if level <= 2 else DARK
        run.font.bold = True
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)


def add_para(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.color.rgb = DARK
        rest = text[len(bold_prefix) :]
        p.add_run(rest)
    else:
        p.add_run(text)
    for run in p.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10.5)
        if run.font.color.rgb is None:
            run.font.color.rgb = DARK


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(item)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10)
        run.font.color.rgb = DARK


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(item)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10)
        run.font.color.rgb = DARK


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_border(cell, "D6DEE8")
    set_cell_margins(cell, 140, 180, 140, 180)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(10.5)
    r.font.color.rgb = BLUE
    p.add_run("\n" + body)
    for run in p.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        if run.font.size is None:
            run.font.size = Pt(10)
    doc.add_paragraph()


def set_doc_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = DARK
    normal.paragraph_format.line_spacing = 1.2
    normal.paragraph_format.space_after = Pt(6)


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("AI 留资窗口识别项目立项文档")
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("面向找房租房 IM 场景的 Agent Workflow 建设方案")
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(11)
    r.font.color.rgb = MUTED

    meta = doc.add_table(rows=3, cols=2)
    style_table(meta, header=False)
    rows = [
        ("项目阶段", "一期已完成 App-服务端-LLM 留资窗口识别链路，进入立项评审与能力扩展阶段"),
        ("核心目标", "提升 IM 场景下经纪人识别留资窗口、引导用户留资的效率和成功率"),
        ("规划方向", "成功留资对话知识库、智能回复 API 接入、留资卡片与物料动作推荐"),
    ]
    for row, (k, v) in zip(meta.rows, rows):
        row.cells[0].text = k
        row.cells[1].text = v
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Microsoft YaHei"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    run.font.size = Pt(9.5)
    doc.add_paragraph()


def add_table_from_rows(doc: Document, rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    style_table(table)
    for i, row_values in enumerate(rows):
        for j, value in enumerate(row_values):
            cell = table.cell(i, j)
            cell.text = value
            if widths:
                cell.width = Cm(widths[j])
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Microsoft YaHei"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    run.font.size = Pt(9)
    doc.add_paragraph()


def build() -> None:
    doc = Document()
    set_doc_styles(doc)
    add_title(doc)

    add_heading(doc, "1. 项目立项背景", 1)
    add_para(
        doc,
        "我爱我家找房租房 App 的核心转化链路，是用户在 C 端浏览房源后进入 IM，与 B 端经纪人沟通，最终留下手机号并进入电话沟通、约看、成交等后续链路。手机号留资是 IM 咨询从线上兴趣转为线下服务的关键节点。",
    )
    add_para(
        doc,
        "当前留资动作主要依赖经纪人经验判断：什么时候问电话、用什么理由问、是否先处理用户顾虑、是否应该发送卡片或物料。这类判断具有强场景性和强时机性，单靠人工经验容易出现不稳定。",
    )
    add_table_from_rows(
        doc,
        [
            ["现状问题", "具体表现", "业务影响"],
            ["时机不稳定", "部分经纪人过早索要手机号，用户还未建立信任或问题未被回答", "用户沉默、拒绝、流失风险增加"],
            ["机会易错过", "用户已表达看房或强意向，但经纪人未及时承接", "高意向线索未被有效转化"],
            ["上下文利用不足", "画像、浏览轨迹、房源匹配度和对话内容分散在不同系统", "经纪人难以及时形成完整判断"],
            ["话术质量不均", "不同经纪人表达差异大，留资理由不够自然", "影响用户信任和平台服务一致性"],
        ],
        [3.0, 7.0, 6.0],
    )

    add_heading(doc, "2. 从业务问题推导项目价值", 1)
    add_para(doc, "本项目的价值不是简单“让 AI 生成一句回复”，而是通过 Agent Workflow 建立一套可持续判断、可解释、可沉淀、可优化的留资决策能力。完整推导链如下：")
    add_numbered(
        doc,
        [
            "业务增长依赖 IM 留资转化：用户在 IM 中留下手机号后，平台和经纪人才有机会推进电话沟通、约看和成交。",
            "留资转化依赖合适窗口期：用户是否愿意留手机号，与其需求明确程度、信任程度、行动意愿和经纪人承接方式相关。",
            "窗口期判断依赖多源上下文：仅看单句聊天不足以判断时机，需要结合画像、浏览轨迹、房源上下文和多轮对话。",
            "多源上下文需要 Agent 记忆和实时分析：Agent 按会话累计上下文，持续识别阶段、意图、风险和留资分。",
            "识别结果需要转化为经纪人可执行动作：B 端实时提示经纪人当前是否适合留资、为什么适合、应该如何承接。",
            "长期价值来自数据沉淀：成功留资案例、经纪人采纳行为和用户反馈可反哺知识库、策略和模型。",
        ],
    )
    add_callout(
        doc,
        "核心价值结论",
        "项目将“经纪人凭经验判断要电话时机”升级为“基于用户画像、行为轨迹、房源上下文和 IM 对话的实时智能决策”，从而提升留资转化效率，并降低不合适时机打扰用户的风险。",
    )

    add_heading(doc, "3. 项目价值", 1)
    add_table_from_rows(
        doc,
        [
            ["价值对象", "价值说明", "衡量指标"],
            ["用户", "减少过早、强硬索要手机号的体验；在确有约看、确认、推荐需求时获得更自然的服务承接", "拒绝率、沉默率、投诉率"],
            ["经纪人", "实时获得当前用户阶段、留资分、风险点和建议策略，降低经验门槛", "提示采纳率、平均响应时长、留资成功率"],
            ["平台", "将 IM 留资经验沉淀为可复用能力，形成策略闭环和数据资产", "IM 留资率、提示后 N 分钟留资率、带看转化率"],
            ["运营/算法", "通过可视化和结构化数据观察不同场景的转化规律，为后续知识库和模型优化提供基础", "样本标注效率、策略命中率、模型迭代周期"],
        ],
        [3.0, 8.0, 5.0],
    )

    add_heading(doc, "4. 当前建设现状", 1)
    add_para(doc, "当前已通过 App-服务端-LLM 的方式完成第一阶段链路验证：用户行为和对话历史可以上传到服务端，服务端累计上下文并调用 LLM 分析，当前已经能够成功提示留资窗口。")
    add_bullets(
        doc,
        [
            "已支持用户行为、用户画像、房源上下文、用户与经纪人对话的累计上传。",
            "已支持按 conversation_id 维护 Agent 上下文记忆。",
            "已支持调用 LLM 识别当前阶段、留资分、置信度、提示等级和判断原因。",
            "已支持 B 端通过流式订阅获取当前会话窗口期结果。",
            "已支持 H5 监控页面查看会话状态、用户画像、聊天记录和当前窗口期判断。",
        ],
    )

    add_heading(doc, "5. 完整逻辑链与工作流", 1)
    add_para(doc, "完整工作流分为上下文采集、Agent 记忆、窗口识别、策略输出、B 端承接、结果回流六个环节。")
    add_table_from_rows(
        doc,
        [
            ["环节", "输入", "处理", "输出"],
            ["上下文采集", "C 端行为、画像、房源、IM 消息", "业务服务端按 conversation_id 上传增量数据", "标准化上下文事件"],
            ["Agent 记忆", "增量上下文事件", "合并画像和房源，追加行为和消息，维护会话状态", "完整会话上下文"],
            ["窗口识别", "完整会话上下文", "LLM 理解意图，规则计算留资分和风险", "阶段、分数、置信度、风险标记"],
            ["策略输出", "窗口识别结果", "判断是否提示、提示等级、建议沟通策略", "B 端可消费的结构化结果"],
            ["B 端承接", "结构化结果和建议话术", "展示轻提示、强提示或风险提醒，经纪人确认后发送", "经纪人实际沟通动作"],
            ["结果回流", "采纳、修改、忽略、是否留资", "记录反馈并用于后续优化", "策略和知识库优化数据"],
        ],
        [2.5, 4.5, 6.0, 4.2],
    )

    add_heading(doc, "6. 前后端交互时序", 1)
    add_para(doc, "前后端交互采用“增量上传 + 流式订阅”的模式：C 端和 IM 服务持续上传上下文，B 端进入聊天页后订阅当前会话的窗口期判断结果。")
    add_table_from_rows(
        doc,
        [
            ["步骤", "发起方", "接收方", "说明"],
            ["1", "B 端 App", "Agent 服务", "建立 SSE 订阅：/agent/window/stream?conversation_id=xxx"],
            ["2", "C 端 App/IM 服务", "Agent 服务", "上传用户画像、行为轨迹、房源上下文"],
            ["3", "IM 服务", "Agent 服务", "上传用户或经纪人的新增聊天消息"],
            ["4", "Agent 服务", "LLM 服务", "携带累计上下文请求阶段和窗口期分析"],
            ["5", "LLM 服务", "Agent 服务", "返回语义判断、意图、风险和策略建议"],
            ["6", "Agent 服务", "Agent 服务", "规则兜底、留资分计算、风险拦截、版本校验"],
            ["7", "Agent 服务", "B 端 App", "通过 SSE 推送最新 analysis 事件"],
            ["8", "B 端 App", "经纪人", "展示留资提示卡、轻提示或风险提醒"],
        ],
        [1.5, 3.0, 3.0, 9.0],
    )
    add_callout(
        doc,
        "时序图说明",
        "在系统设计上，B 端不需要轮询分析结果，而是持续订阅当前 conversation_id。每次 C 端行为或 IM 消息更新后，Agent 会重新分析并主动推送最新状态。",
    )

    add_heading(doc, "7. 本期建设范围", 1)
    add_bullets(
        doc,
        [
            "建立 Agent 上下文记忆：累计用户画像、浏览轨迹、房源上下文和对话历史。",
            "建立留资窗口识别能力：识别用户阶段、计算留资分、输出是否提示经纪人。",
            "建立 B 端实时接收能力：通过流式传输获取最新窗口期判断。",
            "建立可视化监控页面：支持按会话查看当前状态、画像、聊天记录和分析结果。",
            "建立基础风控规则：用户拒绝、费用顾虑未处理、流失风险等场景不触发强留资提示。",
        ],
    )

    add_heading(doc, "8. 未来规划", 1)
    add_heading(doc, "8.1 成功留资对话知识库", 2)
    add_para(doc, "准备一批经纪人与用户成功留资的对话数据，经过清洗和结构化标注后形成知识库。知识库用于增强 AI 在不同场景下生成回复时的规范性、自然度和业务贴合度。")
    add_table_from_rows(
        doc,
        [
            ["知识库字段", "说明"],
            ["scene", "约看前留资、确认房源状态留资、推荐房源留资、处理顾虑后留资等"],
            ["trigger_signal", "用户询问看房、询问是否还在、要求推荐、明确预算和入住时间等"],
            ["conversation_excerpt", "成功留资前 3-5 轮关键对话"],
            ["successful_reply", "经纪人成功引导留资的话术"],
            ["why_it_worked", "成功原因，例如先回应需求、留资理由自然、低压表达"],
            ["do_not_say", "禁忌表达，例如不留电话不能看、先留电话再说"],
        ],
        [4.0, 12.0],
    )

    add_heading(doc, "8.2 留资卡片与物料动作推荐", 2)
    add_para(doc, "将现有留资卡片、房源报告、户型报告等物料提供给 AI，由 Agent 判断当前是否适合推荐经纪人发送工具、钩子卡片或报告物料。")
    add_table_from_rows(
        doc,
        [
            ["推荐动作", "适合时机", "价值"],
            ["留资卡片", "用户明确看房、需要确认房源状态或后续同步", "将手机号收集动作产品化、低压化"],
            ["约看卡片", "用户提出今天、明天、周末等具体看房时间", "承接行动意图，提高约看效率"],
            ["房源报告", "用户关注真实性、价格、周边、配置", "增强信任，减少反复解释"],
            ["户型报告", "用户关注面积、朝向、采光、布局", "用结构化物料回应细节问题"],
            ["相似房源卡片", "用户嫌贵、想比较、问附近还有没有", "从单房源兴趣转向多房源推荐"],
        ],
        [3.5, 6.0, 6.0],
    )

    add_heading(doc, "8.3 接入智能回复生成 API", 2)
    add_para(doc, "公司现有职能部门已具备生成回复的 API。后续工作流中，Agent 负责判断窗口期、提供策略和上下文；智能回复 API 负责生成最终经纪人回复。")
    add_callout(
        doc,
        "未来生成链路",
        "Agent 识别窗口期 -> 检索成功案例知识库 -> 判断是否推荐卡片/物料 -> 调用智能回复生成 API -> Agent 做风控校验 -> 推送给 B 端经纪人确认发送。",
    )

    add_heading(doc, "9. 风险与应对", 1)
    add_table_from_rows(
        doc,
        [
            ["风险", "应对策略"],
            ["AI 误判窗口期", "引入规则兜底、置信度阈值和经纪人确认机制"],
            ["过早提示留资", "用户顾虑未处理时拦截，拒绝后进入冷却"],
            ["提示过多打扰经纪人", "同一会话增加强提示频控，阶段变化时再提示"],
            ["上下文过长影响模型效率", "保留最近消息，长会话生成摘要和需求槽位"],
            ["旧分析覆盖新上下文", "增加 version 控制，只推送最新上下文结果"],
            ["话术不符合平台规范", "接入智能回复 API 和话术风控校验"],
        ],
        [5.0, 11.0],
    )

    add_heading(doc, "10. 验收指标", 1)
    add_table_from_rows(
        doc,
        [
            ["类型", "指标"],
            ["功能指标", "上下文可正确累计；B 端可实时接收分析结果；H5 可查看画像、聊天和窗口期状态"],
            ["业务指标", "IM 留资率提升；提示后 N 分钟内留资率提升；留资后带看转化提升"],
            ["体验指标", "用户拒绝率、沉默率、投诉率不升高；经纪人采纳率提升"],
            ["模型指标", "窗口期判断准确率、误触发率、漏触发率、阶段识别准确率"],
        ],
        [3.5, 12.5],
    )

    add_heading(doc, "11. 总结", 1)
    add_para(
        doc,
        "本项目的立项价值在于，将 IM 留资这一高度依赖经纪人经验的关键转化动作，升级为基于用户画像、行为轨迹、房源上下文和多轮对话的 Agent Workflow。当前链路已经验证能够成功提示留资窗口，具备继续产品化和规模化试点的基础。",
    )
    add_para(
        doc,
        "后续通过成功留资对话知识库、智能回复生成 API 和留资卡片/物料动作推荐，系统将从“判断什么时候适合留资”进一步升级为“判断什么时候用什么话术、什么工具推进留资”，最终形成可沉淀、可评估、可持续优化的 IM 转化能力。",
    )

    doc.save(OUT)


if __name__ == "__main__":
    build()
